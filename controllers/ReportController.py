import os
import uuid
import logging
<<<<<<< HEAD
from flask import Blueprint, request, jsonify, send_file, current_app
from werkzeug.utils import secure_filename
from extentions.db import db
from models.reportModel import QuarterlyReport, ReportFile
=======
from io import BytesIO
from xml.sax.saxutils import escape
from flask import Blueprint, request, jsonify, make_response, send_file
from extentions.db import db
from models.reportModel import QuarterlyReport
from models.projectModel import Project
>>>>>>> 194ab3b (feat: reports)
from datetime import datetime

logger = logging.getLogger(__name__)

report_bp = Blueprint('report_bp', __name__)

POINT_FIELDS = [f"point_{i}" for i in range(1, 18)]

<<<<<<< HEAD
# Fayl yükləmə yalnız 4-cü rüb üçün icazəlidir
FOURTH_QUARTER = 4


def _allowed_file(filename):
    if '.' not in filename:
        return False
    ext = filename.rsplit('.', 1)[1].lower()
    return ext in current_app.config['ALLOWED_REPORT_EXTENSIONS']


def _report_files_folder():
    folder = current_app.config['REPORT_FILES_FOLDER']
    os.makedirs(folder, exist_ok=True)
    return folder


def _get_or_create_report(project_code, quarter_number, year):
    report = QuarterlyReport.query.filter_by(
        project_code=project_code,
        quarter_number=quarter_number,
        year=year
    ).first()
    if not report:
        report = QuarterlyReport(
            project_code=project_code,
            quarter_number=quarter_number,
            year=year,
        )
        db.session.add(report)
        db.session.flush()
    return report
=======
POINT_LABELS = {
    "point_1":  "1. Cari rübdə görülmüş elmi işlər",
    "point_2":  "2. Planlaşdırılmış işlərin yerinə yetirilmə dərəcəsi (%)",
    "point_3":  "3. Əldə edilmiş elmi nəticələr, yenilik",
    "point_4":  "4. Tətbiq olunan metod və yanaşmalar",
    "point_5":  "5. Elmi nəşrlər",
    "point_6":  "6. İxtiralar və patentlər",
    "point_7":  "7. Ezamiyyətlər",
    "point_8":  "8. Elmi ekspedisiyalar",
    "point_9":  "9. Digər tədbirlər",
    "point_10": "10. Elmi məruzələr",
    "point_11": "11. Əldə edilmiş cihaz, avadanlıq, qurğu və mal-materiallar",
    "point_12": "12. Yerli həmkarlarla əlaqələr",
    "point_13": "13. Xarici həmkarlarla əlaqələr",
    "point_14": "14. Kadr hazırlığı",
    "point_15": "15. Sərgilərdə iştirak",
    "point_16": "16. Stajirovka və təcrübə mübadiləsi",
    "point_17": "17. Elmi-kütləvi nəşrlər, mediada çıxışlar və s.",
}


def _fetch_reports(project_code, quarter=None, year=None):
    """Return quarterly reports for a project, optionally filtered by quarter/year,
    ordered chronologically (year, then quarter)."""
    query = QuarterlyReport.query.filter_by(project_code=project_code)
    if quarter:
        query = query.filter_by(quarter_number=quarter)
    if year:
        query = query.filter_by(year=year)
    return query.order_by(
        QuarterlyReport.year.asc(),
        QuarterlyReport.quarter_number.asc()
    ).all()
>>>>>>> 194ab3b (feat: reports)


@report_bp.route('/api/reports/save', methods=['POST'])
def save_report():
    try:
        data = request.get_json()

        required = ['project_code', 'quarter_number', 'year']
        for field in required:
            if field not in data:
                return jsonify({"error": f"{field} is required"}), 400

        project_code   = int(data['project_code'])
        quarter_number = int(data['quarter_number'])
        year           = int(data['year'])

        report = QuarterlyReport.query.filter_by(
            project_code=project_code,
            quarter_number=quarter_number,
            year=year
        ).first()

        if report:
            for field in POINT_FIELDS:
                if field in data:
                    setattr(report, field, data[field])
            report.submission_date = datetime.utcnow()
            db.session.commit()
            return jsonify({
                "message": "Hesabat yeniləndi",
                "report": report.serialize(),
                "status_code": 200
            }), 200
        else:
            new_report = QuarterlyReport(
                project_code=project_code,
                quarter_number=quarter_number,
                year=year,
                **{field: data.get(field) for field in POINT_FIELDS}
            )
            db.session.add(new_report)
            db.session.commit()
            return jsonify({
                "message": "Hesabat yaradıldı",
                "report": new_report.serialize(),
                "status_code": 201
            }), 201

    except Exception as e:
        db.session.rollback()
        logger.error(f"Error saving report: {str(e)}", exc_info=True)
        return jsonify({"error": str(e)}), 500


@report_bp.route('/api/reports/<int:project_code>/<int:quarter_number>/<int:year>', methods=['GET'])
def get_report(project_code, quarter_number, year):
    try:
        report = QuarterlyReport.query.filter_by(
            project_code=project_code,
            quarter_number=quarter_number,
            year=year
        ).first()

        if not report:
            return jsonify({"message": "Hesabat tapılmadı"}), 404

        return jsonify({
            "message": "Hesabat tapıldı",
            "report": report.serialize(),
            "status_code": 200
        }), 200

    except Exception as e:
        logger.error(f"Error getting report: {str(e)}", exc_info=True)
        return jsonify({"error": str(e)}), 500


<<<<<<< HEAD
@report_bp.route('/api/reports/files/upload', methods=['POST'])
def upload_report_files():
    """4-cü rüb hesabatına bir və ya bir neçə fayl (pdf, doc, docx) əlavə edir."""
    try:
        project_code   = request.form.get('project_code')
        quarter_number = request.form.get('quarter_number')
        year           = request.form.get('year')

        if not all([project_code, quarter_number, year]):
            return jsonify({"error": "project_code, quarter_number və year tələb olunur"}), 400

        project_code   = int(project_code)
        quarter_number = int(quarter_number)
        year           = int(year)

        if quarter_number != FOURTH_QUARTER:
            return jsonify({"error": "Fayl yükləmə yalnız 4-cü rüb üçün mümkündür"}), 400

        files = request.files.getlist('files')
        if not files or all(f.filename == '' for f in files):
            return jsonify({"error": "Heç bir fayl seçilməyib"}), 400

        max_size = current_app.config['MAX_REPORT_FILE_SIZE']
        folder = _report_files_folder()

        report = _get_or_create_report(project_code, quarter_number, year)

        saved = []
        for file in files:
            if not file or file.filename == '':
                continue

            if not _allowed_file(file.filename):
                db.session.rollback()
                return jsonify({
                    "error": f"'{file.filename}' faylının tipinə icazə verilmir. "
                             f"Yalnız: {', '.join(sorted(current_app.config['ALLOWED_REPORT_EXTENSIONS']))}"
                }), 400

            # Fayl ölçüsünü yoxla
            file.seek(0, os.SEEK_END)
            size = file.tell()
            file.seek(0)
            if size > max_size:
                db.session.rollback()
                return jsonify({
                    "error": f"'{file.filename}' faylı çox böyükdür (maks. "
                             f"{max_size // (1024 * 1024)} MB)"
                }), 400

            original_name = secure_filename(file.filename) or 'file'
            ext = original_name.rsplit('.', 1)[1].lower()
            stored_name = f"{uuid.uuid4().hex}.{ext}"
            file.save(os.path.join(folder, stored_name))

            report_file = ReportFile(
                report_id=report.id,
                original_filename=original_name,
                stored_filename=stored_name,
                content_type=file.mimetype,
                file_size=size,
            )
            db.session.add(report_file)
            saved.append(report_file)

        db.session.commit()

        return jsonify({
            "message": f"{len(saved)} fayl uğurla yükləndi",
            "files": [f.serialize() for f in saved],
            "status_code": 201
        }), 201

    except Exception as e:
        db.session.rollback()
        logger.error(f"Error uploading report files: {str(e)}", exc_info=True)
        return jsonify({"error": str(e)}), 500


@report_bp.route('/api/reports/files/<int:project_code>/<int:quarter_number>/<int:year>', methods=['GET'])
def list_report_files(project_code, quarter_number, year):
    """Verilmiş hesabata aid faylların siyahısını qaytarır."""
    try:
        report = QuarterlyReport.query.filter_by(
            project_code=project_code,
            quarter_number=quarter_number,
            year=year
        ).first()

        if not report:
            return jsonify({"files": [], "status_code": 200}), 200

        return jsonify({
            "files": [f.serialize() for f in report.files],
            "status_code": 200
        }), 200

    except Exception as e:
        logger.error(f"Error listing report files: {str(e)}", exc_info=True)
        return jsonify({"error": str(e)}), 500


@report_bp.route('/api/reports/files/download/<int:file_id>', methods=['GET'])
def download_report_file(file_id):
    """Faylı yükləmək üçün qaytarır."""
    try:
        report_file = ReportFile.query.get(file_id)
        if not report_file:
            return jsonify({"error": "Fayl tapılmadı"}), 404

        path = os.path.join(current_app.config['REPORT_FILES_FOLDER'], report_file.stored_filename)
        if not os.path.exists(path):
            return jsonify({"error": "Fayl serverdə tapılmadı"}), 404

        return send_file(
            path,
            as_attachment=True,
            download_name=report_file.original_filename,
            mimetype=report_file.content_type or 'application/octet-stream',
        )

    except Exception as e:
        logger.error(f"Error downloading report file: {str(e)}", exc_info=True)
        return jsonify({"error": str(e)}), 500


@report_bp.route('/api/reports/files/<int:file_id>', methods=['DELETE'])
def delete_report_file(file_id):
    """Yüklənmiş faylı silir."""
    try:
        report_file = ReportFile.query.get(file_id)
        if not report_file:
            return jsonify({"error": "Fayl tapılmadı"}), 404

        path = os.path.join(current_app.config['REPORT_FILES_FOLDER'], report_file.stored_filename)
        if os.path.exists(path):
            os.remove(path)

        db.session.delete(report_file)
        db.session.commit()

        return jsonify({"message": "Fayl silindi", "status_code": 200}), 200

    except Exception as e:
        db.session.rollback()
        logger.error(f"Error deleting report file: {str(e)}", exc_info=True)
        return jsonify({"error": str(e)}), 500
=======
@report_bp.route('/api/reports/project/<int:project_code>', methods=['GET'])
def list_reports(project_code):
    """List every quarterly report saved for a project (used by the admin view)."""
    try:
        reports = _fetch_reports(project_code)
        return jsonify({
            "message": "Hesabatlar tapıldı",
            "reports": [r.serialize() for r in reports],
            "status_code": 200
        }), 200
    except Exception as e:
        logger.error(f"Error listing reports: {str(e)}", exc_info=True)
        return jsonify({"error": str(e)}), 500


@report_bp.route('/api/reports-pdf/<int:project_code>', methods=['GET'])
def reports_pdf(project_code):
    """Export a project's quarterly reports as a PDF. Optional ?quarter=&year= filters."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus.flowables import HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase import pdfmetrics

    font_name = "NotoSans"
    try:
        pdfmetrics.registerFont(TTFont(font_name, "./utils/noto_sans/static/NotoSans-Regular.ttf"))
    except Exception as e:
        return {"status": 500, "message": f"Failed to register local font: {str(e)}"}, 500

    quarter = request.args.get('quarter', type=int)
    year = request.args.get('year', type=int)

    reports = _fetch_reports(project_code, quarter, year)
    if not reports:
        return {"status": 404, "message": "Hesabat tapılmadı"}, 404

    project = Project.query.filter_by(project_code=project_code).first()
    project_name = project.project_name if project else ""

    def _rich(value):
        return escape(value if value else "—").replace("\n", "<br/>")

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("ReportTitle", parent=styles["Title"], fontName=font_name,
                                 fontSize=16, alignment=1, spaceAfter=8)
    subtitle_style = ParagraphStyle("ReportSubtitle", fontName=font_name, fontSize=11,
                                    alignment=1, leading=15, spaceAfter=14,
                                    textColor=colors.HexColor("#555555"))
    section_style = ParagraphStyle("ReportSection", fontName=font_name, fontSize=13,
                                   alignment=1, spaceBefore=14, spaceAfter=10,
                                   textColor=colors.HexColor("#1f2937"))
    label_style = ParagraphStyle("ReportLabel", fontName=font_name, fontSize=10, leading=13,
                                 spaceBefore=6, spaceAfter=2,
                                 textColor=colors.HexColor("#465fff"))
    value_style = ParagraphStyle("ReportValue", fontName=font_name, fontSize=11, leading=15,
                                 spaceAfter=4, wordWrap='CJK')

    elements = [Paragraph("Rüblük Elmi-Texniki Hesabat", title_style)]
    subtitle = f"Layihə kodu: {project_code}"
    if project_name:
        subtitle += f"<br/>{escape(project_name)}"
    elements.append(Paragraph(subtitle, subtitle_style))
    elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#d1d5db"),
                               spaceBefore=4, spaceAfter=10))

    for idx, report in enumerate(reports):
        elements.append(Paragraph(f"{report.year} - {report.quarter_number}-ci rüb", section_style))
        for key, label in POINT_LABELS.items():
            elements.append(Paragraph(escape(label), label_style))
            elements.append(Paragraph(_rich(getattr(report, key)), value_style))
        if idx != len(reports) - 1:
            elements.append(PageBreak())

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    doc.build(elements)
    pdf_data = buffer.getvalue()
    buffer.close()

    response = make_response(pdf_data)
    response.headers["Content-Type"] = "application/pdf"
    response.headers["Content-Disposition"] = f"attachment; filename=project_{project_code}_reports.pdf"
    return response


@report_bp.route('/api/reports-docx/<int:project_code>', methods=['GET'])
def reports_docx(project_code):
    """Export a project's quarterly reports as a DOCX. Optional ?quarter=&year= filters."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    quarter = request.args.get('quarter', type=int)
    year = request.args.get('year', type=int)

    reports = _fetch_reports(project_code, quarter, year)
    if not reports:
        return {"status": 404, "message": "Hesabat tapılmadı"}, 404

    project = Project.query.filter_by(project_code=project_code).first()
    project_name = project.project_name if project else ""

    document = Document()

    title = document.add_heading("Rüblük Elmi-Texniki Hesabat", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Layihə kodu: {project_code}")
    if project_name:
        subtitle.add_run("\n" + project_name)

    for idx, report in enumerate(reports):
        heading = document.add_heading(f"{report.year} - {report.quarter_number}-ci rüb", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for key, label in POINT_LABELS.items():
            label_para = document.add_paragraph()
            label_run = label_para.add_run(label)
            label_run.bold = True
            label_run.font.size = Pt(10)
            label_run.font.color.rgb = RGBColor(0x46, 0x5f, 0xff)
            document.add_paragraph(getattr(report, key) or "—")
        if idx != len(reports) - 1:
            document.add_page_break()

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return send_file(
        output,
        as_attachment=True,
        download_name=f"project_{project_code}_reports.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
>>>>>>> 194ab3b (feat: reports)
